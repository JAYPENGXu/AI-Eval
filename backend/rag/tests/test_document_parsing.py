from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings

from rag.chunkers import ChunkOptions, get_chunker
from rag.document_parsing.ir import BlockIR, DocumentIR, PageIR
from rag.document_parsing.parsers import DocxParser, PdfParser, TextParser
from rag.document_parsing.service import execute_parse_run
from rag.document_parsing.validation import DocumentValidationError, validate_document_file
from rag.models import Document, DocumentParseRun, KnowledgeBase


def make_pdf(*, blank_second_page=False, encrypted=False) -> bytes:
    import fitz

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "RAGPilot document parsing architecture and observability " * 3)
    if blank_second_page:
        pdf.new_page()
    options = {}
    if encrypted:
        options = {
            "encryption": fitz.PDF_ENCRYPT_AES_256,
            "owner_pw": "owner-password",
            "user_pw": "user-password",
        }
    data = pdf.tobytes(**options)
    pdf.close()
    return data


def test_markdown_parser_preserves_heading_path(tmp_path):
    path = tmp_path / "guide.md"
    path.write_text("# 系统设计\n\n正文。\n\n## 索引\n\n索引说明。", encoding="utf-8")

    parsed = TextParser().parse(path, mime_type="text/markdown", title=path.name)

    assert parsed.pages[0].blocks[-1].heading_path == ["系统设计", "索引"]
    assert parsed.pages[0].blocks[-1].text == "索引说明。"


def test_docx_parser_preserves_heading_and_table(tmp_path):
    from docx import Document as DocxDocument

    path = tmp_path / "guide.docx"
    docx = DocxDocument()
    docx.add_heading("部署说明", level=1)
    docx.add_paragraph("准备运行环境。")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "参数"
    table.cell(0, 1).text = "值"
    docx.save(path)

    parsed = DocxParser().parse(
        path,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        title=path.name,
    )

    assert any(block.type == "heading" for block in parsed.pages[0].blocks)
    assert any(block.type == "table" and "参数 | 值" in block.text for block in parsed.pages[0].blocks)
    assert parsed.pages[0].blocks[-1].heading_path == ["部署说明"]


def test_pdf_parser_keeps_page_and_bbox(tmp_path):
    path = tmp_path / "guide.pdf"
    path.write_bytes(make_pdf())

    parsed = PdfParser().parse(path, mime_type="application/pdf", title=path.name)

    assert parsed.pages[0].page_number == 1
    assert parsed.pages[0].blocks[0].bbox
    assert "RAGPilot" in parsed.pages[0].text


def test_validation_rejects_disguised_and_encrypted_files():
    disguised = SimpleUploadedFile("fake.pdf", b"not a pdf", content_type="application/pdf")
    with pytest.raises(DocumentValidationError, match="签名"):
        validate_document_file(disguised)

    encrypted = SimpleUploadedFile("locked.pdf", make_pdf(encrypted=True), content_type="application/pdf")
    with pytest.raises(DocumentValidationError, match="加密"):
        validate_document_file(encrypted)


@pytest.mark.parametrize("method", ["token", "sentence", "sentence_window", "semantic", "markdown"])
def test_all_chunkers_keep_ir_provenance(method):
    document = DocumentIR(
        title="guide.pdf",
        mime_type="application/pdf",
        parser="pymupdf",
        parser_version="1",
        metadata={"parse_run_id": 9},
        pages=[
            PageIR(page_number=3, blocks=[
                BlockIR(
                    id="p3-b1",
                    type="paragraph",
                    text="第一段包含检索说明。第二句继续解释。",
                    page=3,
                    heading_path=["检索", "向量召回"],
                    metadata={"paragraph": 4},
                )
            ])
        ],
    )

    chunks = get_chunker(method).split(document, ChunkOptions(chunk_size=100, chunk_overlap=10))

    assert chunks
    assert chunks[0].metadata["page_start"] == 3
    assert chunks[0].metadata["heading_path"] == ["检索", "向量召回"]
    assert chunks[0].metadata["block_ids"] == ["p3-b1"]
    assert chunks[0].metadata["parse_run_id"] == 9


@pytest.mark.django_db
def test_mixed_pdf_only_replaces_scanned_pages_with_ocr(tmp_path, settings):
    settings.MEDIA_ROOT = tmp_path / "media"
    settings.DOCUMENT_NONBLANK_PAGE_MIN_CHARS = 10
    settings.PDF_NATIVE_TEXT_MIN_CHARS = 40
    user = User.objects.create_user(username="parser-owner", password="secret")
    kb = KnowledgeBase.objects.create(owner=user, name="KB")
    document = Document.objects.create(
        kb=kb,
        filename="mixed.pdf",
        file=ContentFile(make_pdf(blank_second_page=True), name="mixed.pdf"),
        file_type="pdf",
        mime_type="application/pdf",
        size_bytes=100,
        status="queued",
    )
    run = DocumentParseRun.objects.create(document=document, status="queued")

    def fake_parse(_self, _path, progress_callback=None, job_callback=None):
        if job_callback:
            job_callback("job-1")
        if progress_callback:
            progress_callback(1, 1)
        return ["# 扫描附件\n\n这是通过 OCR 恢复的扫描页正文，包含足够的有效文字。"]

    with patch("rag.document_parsing.service.PaddleOcrClient.parse", fake_parse):
        execute_parse_run(run.id)

    run.refresh_from_db()
    pages = list(run.pages.order_by("page_number"))
    assert run.status == "completed"
    assert [page.extraction_method for page in pages] == ["native", "ocr"]
    assert pages[1].page_number == 2
    assert pages[1].blocks[0]["heading_path"] == ["扫描附件"]
    assert run.quality_metrics["ocr_page_rate"] == 0.5


@pytest.mark.django_db
def test_older_parse_run_cannot_overwrite_newer_result(tmp_path, settings):
    settings.MEDIA_ROOT = tmp_path / "media"
    user = User.objects.create_user(username="run-owner", password="secret")
    kb = KnowledgeBase.objects.create(owner=user, name="KB")
    document = Document.objects.create(
        kb=kb,
        filename="notes.txt",
        file=ContentFile(b"old parse content", name="notes.txt"),
        file_type="txt",
        mime_type="text/plain",
        size_bytes=17,
        status="queued",
    )
    older = DocumentParseRun.objects.create(document=document, status="queued")
    newer = DocumentParseRun.objects.create(document=document, status="completed")

    execute_parse_run(older.id)

    older.refresh_from_db()
    newer.refresh_from_db()
    assert older.status == "superseded"
    assert newer.status == "completed"


@pytest.mark.django_db
def test_failed_reindex_keeps_previous_chunks(tmp_path, settings):
    from rag.indexing import index_document
    from rag.models import Chunk, DocumentPage

    settings.MEDIA_ROOT = tmp_path / "media"
    user = User.objects.create_user(username="index-owner", password="secret")
    kb = KnowledgeBase.objects.create(owner=user, name="KB")
    document = Document.objects.create(
        kb=kb,
        filename="notes.txt",
        file=ContentFile(b"new index content", name="notes.txt"),
        file_type="txt",
        mime_type="text/plain",
        size_bytes=17,
        status="indexed",
    )
    run = DocumentParseRun.objects.create(document=document, status="completed", parser="text", parser_version="1")
    DocumentPage.objects.create(
        parse_run=run,
        page_number=1,
        text="新的解析内容用于重新索引。",
        blocks=[{
            "id": "p1-b1", "type": "paragraph", "text": "新的解析内容用于重新索引。",
            "page": 1, "heading_path": [], "bbox": [], "confidence": None,
            "metadata": {"paragraph": 1},
        }],
        char_count=13,
    )
    old = Chunk.objects.create(
        document=document, kb=kb, index=0, content="旧索引仍然可用。", embedding=[0.1, 0.2]
    )

    class FailingStore:
        restored = False

        def delete_document(self, _document_id):
            return None

        def index_chunks(self, chunks):
            ids = {chunk.id for chunk in chunks}
            if old.id not in ids:
                raise RuntimeError("vector write failed")
            self.restored = True

    store = FailingStore()
    with (
        patch("rag.indexing.embed_texts", return_value=[[0.3, 0.4]]),
        patch("rag.indexing.get_vector_store", return_value=store),
        pytest.raises(RuntimeError, match="旧索引保持可用"),
    ):
        index_document(document, "sentence", {}, run.id)

    assert document.chunks.count() == 1
    assert document.chunks.get().id == old.id
    assert store.restored is True


@pytest.mark.django_db
def test_celery_redelivery_uses_task_id_lease(tmp_path, settings):
    settings.MEDIA_ROOT = tmp_path / "media"
    settings.DOCUMENT_NONBLANK_PAGE_MIN_CHARS = 5
    user = User.objects.create_user(username="lease-owner", password="secret")
    kb = KnowledgeBase.objects.create(owner=user, name="KB")
    document = Document.objects.create(
        kb=kb,
        filename="notes.txt",
        file=ContentFile("可恢复的解析任务正文。".encode(), name="notes.txt"),
        file_type="txt",
        mime_type="text/plain",
        status="parsing",
    )
    run = DocumentParseRun.objects.create(
        document=document,
        status="running",
        celery_task_id="task-1",
    )

    execute_parse_run(run.id, task_id="other-task")
    run.refresh_from_db()
    assert run.status == "running"
    assert run.pages.count() == 0

    execute_parse_run(run.id, task_id="task-1")
    run.refresh_from_db()
    assert run.status == "completed"
    assert run.pages.count() == 1
